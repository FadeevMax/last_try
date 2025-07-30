import streamlit as st
import os
import json
import pickle
import tempfile
import shutil
import hashlib
from pathlib import Path
import traceback

# Set page config first
st.set_page_config(
    page_title="GTI SOP Assistant - Unified",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)
GITHUB_REPO = os.getenv('GITHUB_REPO', 'FadeevMax/last_try')
GITHUB_TOKEN = os.getenv('GITHUB_TOKEN', '')
GITHUB_BRANCH = 'image-storage'  # Separate branch for images

def upload_to_github(filename, content, folder="images"):
    """Upload file to GitHub"""
    import base64
    import requests
    
    if not GITHUB_TOKEN:
        return None
        
    # Encode content to base64
    content_base64 = base64.b64encode(content).decode('utf-8')
    
    # GitHub API URL
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{folder}/{filename}"
    
    # Check if file exists
    headers = {
    "Authorization": f"Bearer {GITHUB_TOKEN}",  # Change from "token" to "Bearer"
    "Accept": "application/vnd.github.v3+json"
    }
    
    # Get current file SHA if exists (for update)
    sha = None
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        sha = response.json()['sha']
    
    # Prepare data
    data = {
        "message": f"Upload {filename}",
        "content": content_base64,
        "branch": GITHUB_BRANCH
    }
    
    if sha:
        data["sha"] = sha
    
    # Upload file
    response = requests.put(url, headers=headers, json=data)
    
    if response.status_code in [200, 201]:
        return response.json()['content']['download_url']
    else:
        return None

def get_from_github(filename, folder="images"):
    """Get file from GitHub"""
    import requests
    
    if not GITHUB_TOKEN:
        return None
        
    # Direct raw URL
    url = f"https://raw.githubusercontent.com/{GITHUB_REPO}/{GITHUB_BRANCH}/{folder}/{filename}"
    
    try:
        response = requests.get(url)
        if response.status_code == 200:
            return response.content
        return None
    except:
        return None
# Simple credentials loading
def load_api_keys():
    """Load API keys from environment variables only"""
    openai_key = os.getenv('OPENAI_API_KEY', '')
    gemini_key = os.getenv('GEMINI_API_KEY', '')
    return openai_key, gemini_key

# Initialize session state
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'chunks' not in st.session_state:
    st.session_state.chunks = []
if 'vector_db_ready' not in st.session_state:
    st.session_state.vector_db_ready = False
if 'messages' not in st.session_state:
    st.session_state.messages = []

# CSS
st.markdown("""
<style>
.main-header { 
    font-size: 2.5rem; 
    font-weight: bold; 
    background: linear-gradient(90deg, #1e3c72, #2a5298);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-align: center;
    margin-bottom: 2rem;
}
.step-header {
    font-size: 1.5rem;
    font-weight: bold;
    color: #1e3c72;
    margin: 1rem 0;
}
.status-box {
    padding: 1rem;
    border-radius: 8px;
    margin: 1rem 0;
}
.success { background: #d4edda; border-left: 4px solid #28a745; }
.warning { background: #fff3cd; border-left: 4px solid #ffc107; }
.info { background: #d1ecf1; border-left: 4px solid #17a2b8; }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown('<h1 class="main-header">🚀 GTI SOP Assistant - Unified</h1>', unsafe_allow_html=True)
st.markdown("*All-in-one: Document Processing + Vector DB + Chat Interface*")

# Enhanced document chunker with image extraction
def enhanced_chunk_docx(file_content, chunk_size=800):
    """Enhanced DOCX chunker with complete content extraction"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        import io
        import re
        import tempfile
        
        # Create document from uploaded content
        doc = Document(io.BytesIO(file_content))
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        image_counter = 1
        current_context = {'state': None, 'section': None, 'topic': None}
        
        # Create temp directory for images
        if 'stored_images' not in st.session_state:
          st.session_state.stored_images = {}
        
        def update_context(text):
            """Update current context based on text content"""
            text_upper = text.upper()
            
            # State detection
            state_patterns = {
                'OH': [r'\bOHIO\b', r'\bOH\b(?!\w)'],
                'MD': [r'\bMARYLAND\b', r'\bMD\b(?!\w)'],
                'NJ': [r'\bNEW\s+JERSEY\b', r'\bNJ\b(?!\w)'],
                'IL': [r'\bILLINOIS\b', r'\bIL\b(?!\w)'],
                'NY': [r'\bNEW\s+YORK\b', r'\bNY\b(?!\w)'],
                'NV': [r'\bNEVADA\b', r'\bNV\b(?!\w)'],
                'MA': [r'\bMASSACHUSETTS\b', r'\bMA\b(?!\w)']
            }
            
            for state, patterns in state_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text_upper):
                        current_context['state'] = state
                        if 'RISE' in text_upper:
                            current_context['section'] = 'RISE'
                        elif 'REGULAR' in text_upper:
                            current_context['section'] = 'REGULAR'
                        break
            
            # Topic detection
            if 'PRICING' in text_upper or 'MENU PRICE' in text_upper:
                current_context['topic'] = 'PRICING'
            elif 'BATTER' in text_upper:
                current_context['topic'] = 'BATTERIES'
            elif 'BATCH SUB' in text_upper:
                current_context['topic'] = 'BATCH_SUB'
            elif 'DELIVERY DATE' in text_upper:
                current_context['topic'] = 'DELIVERY_DATE'
            elif 'ORDER LIMIT' in text_upper:
                current_context['topic'] = 'ORDER_LIMIT'
        
        # Enhanced caption pattern matching (based on your code)
        def clean_caption(text):
            """Enhanced text cleaning with better normalization"""
            import unicodedata
            cleaned = unicodedata.normalize('NFKC', text)
            cleaned = re.sub(r"\s+", " ", cleaned).strip()
            cleaned = cleaned.replace("–", "-").replace("—", "-").replace(""", '"').replace(""", '"')
            cleaned = cleaned.replace("'", "'").replace("'", "'")
            # Remove excessive punctuation
            cleaned = re.sub(r'[.]{2,}', '.', cleaned)
            return cleaned

        def extract_label(text):
            """Enhanced label extraction supporting multiple formats"""
            text = clean_caption(text)
            
            # Pattern for "Image X: description" format
            caption_pattern = re.compile(r"^Image\s+(\d+)\s*[:.]?\s*(.*?)(?:\.|$)", re.IGNORECASE)
            figure_pattern = re.compile(r"^Figure\s+(\d+)\s*[:.]?\s*(.*?)(?:\.|$)", re.IGNORECASE)
            
            # Try Image pattern first
            m = caption_pattern.match(text)
            if m:
                idx = int(m.group(1))
                desc = m.group(2).strip().rstrip(".")
                return f"Image {idx}: {desc}" if desc else f"Image {idx}"
            
            # Try Figure pattern
            m = figure_pattern.match(text)
            if m:
                idx = int(m.group(1))
                desc = m.group(2).strip().rstrip(".")
                return f"Figure {idx}: {desc}" if desc else f"Figure {idx}"
            
            # Look for descriptive patterns without numbers
            descriptive_patterns = [
                r'([^.]+\s+example\s*[^.]*)',
                r'([^.]+\s+sheet\s*[^.]*)',
                r'([^.]+\s+form\s*[^.]*)',
                r'([^.]+\s+format\s*[^.]*)',
                r'([^.]+\s+setup\s*[^.]*)',
                r'([^.]+\s+process\s*[^.]*)',
                r'([^.]+\s+workflow\s*[^.]*)',
                r'([^.]+\s+template\s*[^.]*)'
            ]
            
            for pattern in descriptive_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    desc = match.group(1).strip()
                    if len(desc) > 5 and len(desc) < 80:  # Reasonable caption length
                        return desc
            
            return None

        def extract_images_with_enhanced_labels():
            """Extract images with sophisticated caption matching"""
            nonlocal image_counter
            
            # Collect all document items in order
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            
            items = []
            body = doc.element.body
            position = 0
            
            for child in body.iterchildren():
                if isinstance(child, CT_P):
                    para = Paragraph(child, doc)
                    
                    # Check for images in paragraph
                    has_image = False
                    for run in para.runs:
                        if 'graphic' in run._element.xml:
                            for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                    rel_id = blip.get(qn('r:embed'))
                                    if rel_id and rel_id in doc.part.related_parts:
                                        image_part = doc.part.related_parts[rel_id]
                                        items.append({
                                            'type': 'image', 
                                            'content': image_part, 
                                            'position': position,
                                            'paragraph_text': para.text.strip(),
                                            'element': child
                                        })
                                        has_image = True
                    
                    # Add text if it exists and doesn't have images
                    if para.text.strip() and not has_image:
                        items.append({
                            'type': 'text', 
                            'content': para.text.strip(), 
                            'position': position
                        })
                        
                elif isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                # Check for images in table cells
                                for run in para.runs:
                                    if 'graphic' in run._element.xml:
                                        for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                            for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                                rel_id = blip.get(qn('r:embed'))
                                                if rel_id and rel_id in doc.part.related_parts:
                                                    image_part = doc.part.related_parts[rel_id]
                                                    items.append({
                                                        'type': 'image', 
                                                        'content': image_part, 
                                                        'position': position,
                                                        'paragraph_text': para.text.strip(),
                                                        'element': para._element
                                                    })
                
                position += 1

            # Process images with enhanced caption matching
            images = []
            i = 0
            
            while i < len(items):
                if items[i]['type'] == 'image':
                    image_part = items[i]['content']
                    
                    # Look for caption in multiple places
                    label = None
                    
                    # 1. Check if the image's paragraph contains caption text
                    if items[i].get('paragraph_text'):
                        potential_label = extract_label(items[i]['paragraph_text'])
                        if potential_label:
                            label = potential_label
                    
                    # 2. Look ahead for following caption
                    if not label:
                        for j in range(i + 1, min(i + 4, len(items))):  # Look ahead up to 3 items
                            if items[j]['type'] == 'text':
                                potential_label = extract_label(items[j]['content'])
                                if potential_label:
                                    label = potential_label
                                    break
                    
                    # 3. Look behind for preceding caption
                    if not label:
                        for j in range(max(0, i - 3), i):  # Look behind up to 3 items
                            if items[j]['type'] == 'text':
                                potential_label = extract_label(items[j]['content'])
                                if potential_label:
                                    label = potential_label
                                    break
                    
                    # Default label if none found
                    if not label:
                        label = f"Image {image_counter}"
                    
                    # Save image file
                    image_extension = image_part.content_type.split('/')[-1]
                    if image_extension == 'jpeg':
                        image_extension = 'jpg'
                    elif image_extension not in ['jpg', 'png', 'gif', 'bmp', 'webp']:
                        image_extension = 'png'
                        
                    image_filename = f"image_{image_counter}.{image_extension}"
                    image_url = upload_to_github(image_filename, image_part.blob)
                    
                    images.append({
                        'filename': image_filename,
                        'url': image_url,  # Store URL instead of path
                        'label': label,
                        'number': image_counter,
                        'position': items[i]['position']
                    })
                    
                    image_counter += 1
                    
                i += 1
            
            return images
        
        def process_paragraph(para):
            """Process paragraph with enhanced extraction"""
            text = para.text.strip()
            images = []
            
            # Extract images from paragraph
            try:
                images = extract_images_from_element(para._element)
            except:
                pass
            
            return text, images
        
        def process_table(table):
            """Extract table content"""
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' | '.join(row_text))
            return '\n'.join(table_text)
        
        # Extract all images with enhanced labels first
        all_images = extract_images_with_enhanced_labels()
        
        # Process document body elements in order
        all_content = []
        position = 0
        
        # Process paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # Paragraph
                para = next(p for p in doc.paragraphs if p._element == element)
                text = para.text.strip()
                if text:
                    all_content.append({
                        'type': 'paragraph',
                        'text': text,
                        'position': position
                    })
            elif element.tag.endswith('tbl'):
                # Table
                table = next(t for t in doc.tables if t._element == element)
                table_text = process_table(table)
                if table_text:
                    all_content.append({
                        'type': 'table',
                        'text': table_text,
                        'position': position
                    })
            position += 1
        
        # Also check headers and footers
        for section in doc.sections:
            # Headers
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_content.insert(0, {
                            'type': 'header',
                            'text': text,
                            'position': -1  # Headers come first
                        })
            
            # Footers
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        all_content.append({
                            'type': 'footer',
                            'text': text,
                            'position': 999  # Footers come last
                        })
        
        # Create image lookup by position for efficient matching
        image_by_position = {}
        for img in all_images:
            pos = img.get('position', 0)
            if pos not in image_by_position:
                image_by_position[pos] = []
            image_by_position[pos].append(img)
        
        # Debug: Show image extraction results
        if all_images:
            st.success(f"✅ Extracted {len(all_images)} images with enhanced labels")
            with st.expander("🖼️ Image Details"):
                for i, img in enumerate(all_images[:5]):  # Show first 5
                    st.write(f"**{img['label']}** - Position: {img.get('position', 'N/A')}")
        else:
            st.warning("⚠️ No images found in document")
        
        # Create chunks with strict position-based image attachment
        chunk_ranges = []  # Track [start_pos, end_pos] for each chunk
        current_position = 0
        
        for content in all_content:
            text = content['text']
            content_position = content.get('position', current_position)
            
            if not text:
                continue
            
            # Update context
            update_context(text)
            
            # If adding this text would exceed chunk size, save current chunk
            if len(current_chunk) + len(text) > chunk_size and current_chunk:
                # Record the position range for this chunk
                chunk_ranges.append([current_position - 10, current_position])  # Rough range
                
                chunks.append({
                    'chunk_id': chunk_id,
                    'text': current_chunk.strip(),
                    'images': [],  # Will be filled in next step
                    'start_pos': current_position - 10,
                    'end_pos': current_position,
                    'metadata': {
                        'states': [current_context['state']] if current_context['state'] else [],
                        'sections': [current_context['section']] if current_context['section'] else [],
                        'topics': [current_context['topic']] if current_context['topic'] else [],
                        'word_count': len(current_chunk.split()),
                        'has_images': False,
                        'image_count': 0
                    }
                })
                chunk_id += 1
                current_chunk = ""
            
            current_chunk += text + " "
            current_position = content_position
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'chunk_id': chunk_id,
                'text': current_chunk.strip(),
                'images': [],
                'start_pos': current_position - 10,
                'end_pos': current_position + 10,
                'metadata': {
                    'states': [current_context['state']] if current_context['state'] else [],
                    'sections': [current_context['section']] if current_context['section'] else [],
                    'topics': [current_context['topic']] if current_context['topic'] else [],
                    'word_count': len(current_chunk.split()),
                    'has_images': False,
                    'image_count': 0
                }
            })
        
        # Now assign images to chunks based on strict document position adjacency
        for chunk in chunks:
            chunk_images = []
            start_pos = chunk.get('start_pos', 0)
            end_pos = chunk.get('end_pos', 999)
            
            # Find images that are immediately before, within, or after this chunk
            # Check positions: [start-3, start-2, start-1, start...end, end+1, end+2, end+3]
            for check_pos in range(start_pos - 3, end_pos + 4):
                if check_pos in image_by_position:
                    chunk_images.extend(image_by_position[check_pos])
            
            # Sort images by their document position to maintain order
            chunk_images.sort(key=lambda x: x.get('position', 0))
            
            # Update chunk with images
            chunk['images'] = chunk_images
            chunk['metadata']['has_images'] = len(chunk_images) > 0
            chunk['metadata']['image_count'] = len(chunk_images)
            
            # Clean up temporary position fields
            if 'start_pos' in chunk:
                del chunk['start_pos']
            if 'end_pos' in chunk:
                del chunk['end_pos']
        
        # Debug: Show how many images were assigned
        total_assigned = sum(len(chunk['images']) for chunk in chunks)
        st.success(f"📸 Position-based assignment: {total_assigned} images assigned to chunks")
        
        if total_assigned < len(all_images):
            unassigned_count = len(all_images) - total_assigned
            st.warning(f"⚠️ {unassigned_count} images were not adjacent to any text chunks (this is normal for isolated images)")
        
        return chunks
        
    except Exception as e:
        st.error(f"Error processing document: {e}")
        st.code(traceback.format_exc())
        return []

# Enhanced search with context awareness
def enhanced_search(chunks, query, top_k=5):
    """Enhanced search with context and metadata awareness"""
    import re
    
    query_lower = query.lower()
    query_words = set(query_lower.split())
    
    # Detect if this is a comparison/analytical question
    analytical_keywords = [
        'how many', 'what states', 'which state', 'highest', 'lowest', 'compare', 
        'all states', 'total', 'maximum', 'minimum', 'list all', 'differences',
        'across states', 'between states', 'summary', 'overview'
    ]
    
    is_analytical = any(keyword in query_lower for keyword in analytical_keywords)
    
    if is_analytical:
        # For analytical questions, we need broader context
        top_k = min(15, len(chunks))  # Get more chunks for analysis
    
    # Extract query intent
    query_state = None
    query_section = None
    query_topics = []
    
    # State detection
    state_patterns = {
        'OH': [r'\boh\b', r'\bohio\b'],
        'MD': [r'\bmd\b', r'\bmaryland\b'],
        'NJ': [r'\bnj\b', r'\bnew jersey\b', r'\bjersey\b'],
        'IL': [r'\bil\b', r'\billinois\b'],
        'NY': [r'\bny\b', r'\bnew york\b'],
        'NV': [r'\bnv\b', r'\bnevada\b'],
        'MA': [r'\bma\b', r'\bmassachusetts\b']
    }
    
    for state, patterns in state_patterns.items():
        for pattern in patterns:
            if re.search(pattern, query_lower):
                query_state = state
                break
    
    # Section detection
    if 'rise' in query_lower or 'internal' in query_lower:
        query_section = 'RISE'
    elif 'regular' in query_lower or 'wholesale' in query_lower:
        query_section = 'REGULAR'
    
    # Topic detection
    topic_keywords = {
        'PRICING': ['price', 'pricing', 'cost', 'discount', 'menu'],
        'BATTERIES': ['battery', 'batteries', 'separate', 'invoice'],
        'BATCH_SUB': ['batch', 'sub', 'substitution', 'split'],
        'DELIVERY_DATE': ['delivery', 'date', 'schedule'],
        'ORDER_LIMIT': ['limit', 'maximum', 'max', 'unit'],
        'LESS_AVAILABLE': ['less', 'available', 'partial', 'shortage']
    }
    
    for topic, keywords in topic_keywords.items():
        if any(keyword in query_lower for keyword in keywords):
            query_topics.append(topic)
    
    # Score chunks
    scored_chunks = []
    for chunk in chunks:
        score = 0
        text_lower = chunk['text'].lower()
        text_words = set(text_lower.split())
        metadata = chunk.get('metadata', {})
        
        # Base keyword overlap score
        overlap = len(query_words.intersection(text_words))
        score += overlap / max(len(query_words), 1)
        
        # Boost for exact phrase matches
        for word in query_words:
            if len(word) > 3 and word in text_lower:
                score += 0.3
        
        # Boost for metadata matches
        if query_state and query_state in metadata.get('states', []):
            score += 0.5
        
        if query_section and query_section in metadata.get('sections', []):
            score += 0.3
        
        for topic in query_topics:
            if topic in metadata.get('topics', []):
                score += 0.4
        
        # Boost for images if visual content is requested
        if any(word in query_lower for word in ['image', 'show', 'example', 'visual']):
            if metadata.get('has_images'):
                score += 0.3
        
        # Boost for longer, more complete content
        if len(chunk['text']) > 200:
            score += 0.1
        
        if score > 0:
            scored_chunks.append({
                'chunk': chunk,
                'score': score,
                'chunk_id': chunk['chunk_id'],
                'search_types': []
            })
    
    # Sort by score and return top results
    scored_chunks.sort(key=lambda x: x['score'], reverse=True)
    return scored_chunks[:top_k]

# Simple LLM client
class SimpleLLMClient:
    def __init__(self):
        self.openai_key = ""
        self.gemini_key = ""
    
    def setup_keys(self, openai_key, gemini_key):
        self.openai_key = openai_key
        self.gemini_key = gemini_key
    
    def generate_response(self, model, context, temperature=0.1):
        """Generate response using selected model"""
        if model == "Gemini 2.0 Flash" and self.gemini_key:
            return self._generate_gemini(context, temperature)
        elif "GPT" in model and self.openai_key:
            return self._generate_openai(context, model, temperature)
        else:
            return "⚠️ Please configure API keys in the sidebar to use AI models."
    
    def _generate_gemini(self, context, temperature):
        try:
            import requests
            
            headers = {
                'Content-Type': 'application/json',
                'X-goog-api-key': self.gemini_key
            }
            
            prompt = f"""You are a GTI SOP Assistant. Answer based ONLY on the provided documentation.

CONTEXT:
{context}

Provide a clear, specific answer based only on the information above. If the information is not in the context, say so clearly."""
            
            data = {
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {
                    "temperature": temperature,
                    "maxOutputTokens": 1000
                }
            }
            
            response = requests.post(
                "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent",
                headers=headers,
                json=data,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    return result['candidates'][0]['content']['parts'][0]['text']
            
            return f"Error: {response.status_code} - {response.text}"
            
        except Exception as e:
            return f"Error calling Gemini: {str(e)}"
    
    def _generate_openai(self, context, model, temperature):
        try:
            from openai import OpenAI
            
            client = OpenAI(api_key=self.openai_key)
            
            model_map = {
                "GPT-4": "gpt-4",
                "GPT-4 Mini": "gpt-4o-mini"
            }
            
            response = client.chat.completions.create(
                model=model_map.get(model, "gpt-4o-mini"),
                messages=[
                    {
                        "role": "system", 
                        "content": "You are a GTI SOP Assistant. Answer based ONLY on provided documentation."
                    },
                    {"role": "user", "content": context}
                ],
                max_tokens=1000,
                temperature=temperature
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            return f"Error calling OpenAI: {str(e)}"

# Initialize LLM client
if 'llm_client' not in st.session_state:
    st.session_state.llm_client = SimpleLLMClient()

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # API Keys
    with st.expander("🔑 API Keys"):
        openai_key_default, gemini_key_default = load_api_keys()
        
        gemini_key = st.text_input(
            "Gemini API Key", 
            value=gemini_key_default if gemini_key_default else "",
            type="password"
        )
        openai_key = st.text_input(
            "OpenAI API Key", 
            value=openai_key_default if openai_key_default else "",
            type="password"
        )
        
        if st.button("💾 Save Keys"):
            st.session_state.llm_client.setup_keys(openai_key, gemini_key)
            st.success("Keys saved!")
    
    # Model selection
    model = st.selectbox(
        "🤖 AI Model",
        ["Gemini 2.0 Flash", "GPT-4 Mini", "GPT-4"],
        help="Gemini 2.0 Flash is recommended for cost/performance"
    )
    
    temperature = st.slider("🌡️ Temperature", 0.0, 1.0, 0.1, 0.1)
    
    # Status
    st.divider()
    st.header("📊 Status")
    
    if st.session_state.processing_complete:
        st.success(f"✅ Document processed ({len(st.session_state.chunks)} chunks)")
    else:
        st.warning("⏳ No document processed yet")
    
    if st.session_state.vector_db_ready:
        st.success("✅ Search ready")
    else:
        st.info("ℹ️ Upload document to enable search")

# Main interface with tabs
tab1, tab2, tab3 = st.tabs(["📄 Process Document", "🔍 Search & Chat", "📋 View Chunks"])

# Tab 1: Document Processing
with tab1:
    st.markdown('<div class="step-header">Step 1: Upload and Process Document</div>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Upload your DOCX file",
        type=['docx'],
        help="Upload the GTI SOP document for processing"
    )
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        chunk_size = st.number_input("Chunk Size (characters)", 500, 2000, 800)
    
    with col2:
        if st.button("🚀 Process Document", disabled=uploaded_file is None):
            if uploaded_file is not None:
                with st.spinner("Processing document..."):
                    try:
                        # Read file content
                        file_content = uploaded_file.read()
                        
                        # Process document with enhanced extraction
                        chunks = enhanced_chunk_docx(file_content, chunk_size)
                        
                        if chunks:
                            st.session_state.chunks = chunks
                            st.session_state.processing_complete = True
                            st.session_state.vector_db_ready = True
                            
                            # Count chunks with images for debugging
                            chunks_with_images = sum(1 for c in chunks if c.get('images'))
                            total_images = sum(len(c.get('images', [])) for c in chunks)
                            
                            st.success(f"✅ Document processed successfully! Created {len(chunks)} chunks.")
                            st.info(f"📸 {chunks_with_images} chunks contain {total_images} total images")
                            
                            # Show preview
                            with st.expander("📖 Preview First Chunk"):
                                if chunks:
                                    preview_chunk = chunks[0]
                                    st.write(f"**Chunk 0:** {preview_chunk['text'][:200]}...")
                                    st.json(preview_chunk['metadata'])
                        else:
                            st.error("❌ Failed to process document. Please check the file format.")
                            
                    except Exception as e:
                        st.error(f"❌ Error processing document: {str(e)}")
                        st.code(traceback.format_exc())

# Tab 2: Search and Chat
with tab2:
    st.markdown('<div class="step-header">Step 2: Search and Chat</div>', unsafe_allow_html=True)
    
    if not st.session_state.processing_complete:
        st.markdown('<div class="status-box warning">⚠️ Please process a document first in the "Process Document" tab.</div>', unsafe_allow_html=True)
    else:
        # Setup API keys
        st.session_state.llm_client.setup_keys(openai_key, gemini_key)
        
        # Chat interface
        st.markdown("### 💬 Chat with your SOP")
        
        # Display chat history
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.write(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask about GTI procedures..."):
            # Add user message
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            with st.chat_message("user"):
                st.write(prompt)
            
            # Generate response
            with st.chat_message("assistant"):
                with st.spinner("Searching and generating response..."):
                    # Search for relevant chunks with enhanced search
                    search_results = enhanced_search(st.session_state.chunks, prompt, top_k=5)
                    
                    if search_results:
                        # Detect if this is an analytical question requiring cross-chunk analysis
                        query_lower = prompt.lower()
                        analytical_keywords = [
                            'how many', 'what states', 'which state', 'highest', 'lowest', 'compare', 
                            'all states', 'total', 'maximum', 'minimum', 'list all', 'differences',
                            'across states', 'between states', 'summary', 'overview'
                        ]
                        
                        is_analytical = any(keyword in query_lower for keyword in analytical_keywords)
                        
                        if is_analytical:
                            # For analytical questions, build structured context with all data
                            context_parts = [f"USER QUESTION: {prompt}\n\nCOMPREHENSIVE DATA FOR ANALYSIS:"]
                            
                            # Extract structured data from all relevant chunks
                            state_data = {}
                            all_topics = set()
                            
                            for i, result in enumerate(search_results):
                                chunk = result['chunk']
                                metadata = chunk.get('metadata', {})
                                states = metadata.get('states', [])
                                topics = metadata.get('topics', [])
                                
                                context_parts.append(f"\n--- Section {i+1} (Score: {result['score']:.2f}) ---")
                                
                                # Add metadata for analysis
                                if states:
                                    context_parts.append(f"STATES: {', '.join(states)}")
                                if topics:
                                    context_parts.append(f"TOPICS: {', '.join(topics)}")
                                    all_topics.update(topics)
                                
                                # Track state-specific data
                                for state in states:
                                    if state not in state_data:
                                        state_data[state] = []
                                    state_data[state].append({
                                        'text': chunk['text'],
                                        'topics': topics,
                                        'score': result['score']
                                    })
                                
                                context_parts.append(f"CONTENT: {chunk['text']}")
                            
                            # Add analytical summary
                            context_parts.append(f"\n--- SUMMARY FOR ANALYSIS ---")
                            context_parts.append(f"TOTAL SECTIONS FOUND: {len(search_results)}")
                            context_parts.append(f"STATES MENTIONED: {', '.join(sorted(state_data.keys()))}")
                            context_parts.append(f"TOPICS COVERED: {', '.join(sorted(all_topics))}")
                            
                            # Add instruction for analytical response
                            context_parts.append(f"\nINSTRUCTION: Please analyze ALL the provided sections to answer the user's question. Look across all states and data points to provide a comprehensive comparison or summary.")
                            
                        else:
                            # Regular context building for specific questions
                            context_parts = [f"USER QUESTION: {prompt}\n\nRELEVANT DOCUMENTATION:"]
                            
                            for i, result in enumerate(search_results):
                                context_parts.append(f"\n--- Section {i+1} (Score: {result['score']:.2f}) ---")
                                context_parts.append(result['chunk']['text'])
                        
                        context = '\n'.join(context_parts)
                        
                        # Generate answer
                        answer = st.session_state.llm_client.generate_response(model, context, temperature)
                        
                        st.write(answer)
                        
                        # Display images from relevant chunks with smart filtering
                        relevant_images = []
                        query_lower = prompt.lower()
                        query_words = set(query_lower.split())
                        
                        # Extract key concepts from query for better matching
                        concept_keywords = {
                            'order': ['order', 'ordering', 'purchase'],
                            'form': ['form', 'template', 'document'],
                            'invoice': ['invoice', 'billing', 'payment'],
                            'delivery': ['delivery', 'shipping', 'schedule'],
                            'battery': ['battery', 'batteries'],
                            'pricing': ['price', 'pricing', 'cost'],
                            'substitution': ['sub', 'substitution', 'batch'],
                            'rise': ['rise', 'internal'],
                            'regular': ['regular', 'wholesale'],
                            'limit': ['limit', 'maximum', 'max'],
                            'note': ['note', 'notes', 'required'],
                            'split': ['split', 'splitting']
                        }
                        
                        query_concepts = []
                        for concept, keywords in concept_keywords.items():
                            if any(kw in query_lower for kw in keywords):
                                query_concepts.append(concept)
                        
                        # Process images from relevant chunks automatically
                        for result in search_results[:3]:  # Only check top 3 most relevant chunks
                            chunk = result['chunk']
                            if chunk.get('images') and result['score'] > 0.3:  # Only from well-matching chunks
                                for img in chunk['images']:
                                    img_label_lower = img['label'].lower()
                                    img_words = set(img_label_lower.split())
                                    
                                    # Calculate image relevance to the query
                                    relevance_score = 0
                                    
                                    # Direct word overlap (high value)
                                    word_overlap = len(query_words.intersection(img_words))
                                    relevance_score += word_overlap * 2
                                    
                                    # Concept overlap (medium value)
                                    img_concepts = []
                                    for concept, keywords in concept_keywords.items():
                                        if any(kw in img_label_lower for kw in keywords):
                                            img_concepts.append(concept)
                                    
                                    concept_overlap = len(set(query_concepts).intersection(set(img_concepts)))
                                    relevance_score += concept_overlap * 1.5
                                    
                                    # Chunk relevance bonus (the better the chunk matches, the more relevant its images)
                                    chunk_bonus = result['score'] * 1.0
                                    relevance_score += chunk_bonus
                                    
                                    # State/section alignment bonus
                                    chunk_metadata = chunk.get('metadata', {})
                                    if query_concepts:
                                        # If query has specific concepts, image should match them
                                        min_threshold = 1.5
                                    else:
                                        # For general queries, be more lenient with images from top chunks
                                        min_threshold = 1.0
                                    
                                    # Include image if it's relevant enough
                                    if relevance_score >= min_threshold:
                                        relevant_images.append({
                                            'img': img,
                                            'score': relevance_score,
                                            'chunk_score': result['score'],
                                            'chunk_rank': search_results.index(result) + 1
                                        })
                        
                        # Sort by relevance score, then by chunk rank
                        relevant_images.sort(key=lambda x: (x['score'], -x['chunk_rank']), reverse=True)
                        
                        # Debug: Show what images were found
                        total_images = sum(len(r['chunk'].get('images', [])) for r in search_results)
                        st.write(f"🔍 Debug: Found {len(relevant_images)} relevant images (from {total_images} total in chunks)")
                        
                        if relevant_images:
                            st.markdown("### 📸 Related Images")
                            # Show up to 3 most relevant images
                            top_images = relevant_images[:3]
                            
                            if len(top_images) == 1:
                                # Single image - show larger
                                img_data = top_images[0]
                                img = img_data['img']
                                try:
                                    if img.get('url'):
                                        # Use URL directly if available
                                        st.image(img['url'], caption=img['label'], use_container_width=True)
                                    else:
                                        # Fallback to downloading
                                        image_data = get_from_github(img['filename'])
                                        if image_data:
                                            import io
                                            from PIL import Image
                                            image = Image.open(io.BytesIO(image_data))
                                            st.image(image, caption=img['label'], use_container_width=True)
                                except Exception as e:
                                    st.error(f"Cannot display {img['filename']}: {e}")
                            else:
                                # Multiple images - show in columns
                                cols = st.columns(min(3, len(top_images)))
                                for idx, img_data in enumerate(top_images):
                                    img = img_data['img']
                                    with cols[idx]:
                                        try:
                                            if img.get('url'):
                                                # Use URL directly if available
                                                st.image(img['url'], caption=img['label'], use_container_width=True)
                                            else:
                                                # Fallback to downloading
                                                image_data = get_from_github(img['filename'])
                                                if image_data:
                                                    import io
                                                    from PIL import Image
                                                    image = Image.open(io.BytesIO(image_data))
                                                    st.image(image, caption=img['label'], use_container_width=True)
                                        except Exception as e:
                                            st.error(f"Cannot display {img['filename']}: {e}")
                        elif total_images > 0:
                            st.info(f"💡 Found {total_images} images in sections, but none were closely related to your specific query.")
                        
                        # Show search details
                        with st.expander("🔍 Search Details"):
                            st.write(f"Found {len(search_results)} relevant sections:")
                            for i, result in enumerate(search_results):
                                chunk = result['chunk']
                                st.write(f"**Match {i+1}** (Score: {result['score']:.2f})")
                                
                                # Show metadata
                                metadata = chunk.get('metadata', {})
                                if metadata.get('states') or metadata.get('sections') or metadata.get('topics'):
                                    meta_info = []
                                    if metadata.get('states'):
                                        meta_info.append(f"States: {', '.join(metadata['states'])}")
                                    if metadata.get('sections'):
                                        meta_info.append(f"Type: {', '.join(metadata['sections'])}")
                                    if metadata.get('topics'):
                                        meta_info.append(f"Topics: {', '.join(metadata['topics'])}")
                                    if metadata.get('has_images'):
                                        meta_info.append(f"📸 {metadata.get('image_count', 0)} images")
                                    
                                    st.caption(" | ".join(meta_info))
                                
                                st.write(f"Preview: {chunk['text'][:150]}...")
                                st.divider()
                    else:
                        answer = "I couldn't find relevant information for your query. Please try rephrasing your question."
                        st.write(answer)
                    
                    # Save assistant message
                    st.session_state.messages.append({"role": "assistant", "content": answer})
        
        # Clear chat button
        if st.button("🗑️ Clear Chat"):
            st.session_state.messages = []
            st.rerun()

# Tab 3: View Chunks
with tab3:
    st.markdown('<div class="step-header">Step 3: Review Processed Chunks</div>', unsafe_allow_html=True)
    
    if not st.session_state.processing_complete:
        st.markdown('<div class="status-box info">ℹ️ Process a document first to view chunks.</div>', unsafe_allow_html=True)
    else:
        st.write(f"**Total Chunks:** {len(st.session_state.chunks)}")
        
        # Search chunks
        search_term = st.text_input("🔍 Search chunks:", placeholder="Enter keywords to filter chunks...")
        
        chunks_to_show = st.session_state.chunks
        
        if search_term:
            chunks_to_show = [
                chunk for chunk in st.session_state.chunks 
                if search_term.lower() in chunk['text'].lower()
            ]
            st.write(f"**Filtered:** {len(chunks_to_show)} chunks match '{search_term}'")
        
        # Pagination
        chunks_per_page = 5
        total_pages = (len(chunks_to_show) + chunks_per_page - 1) // chunks_per_page
        
        if total_pages > 1:
            page = st.number_input("Page", 1, total_pages, 1) - 1
        else:
            page = 0
        
        start_idx = page * chunks_per_page
        end_idx = min(start_idx + chunks_per_page, len(chunks_to_show))
        
        # Display chunks
        for i in range(start_idx, end_idx):
            chunk = chunks_to_show[i]
            
            with st.expander(f"Chunk {chunk['chunk_id']} - {len(chunk['text'])} chars"):
                st.write("**Text:**")
                st.write(chunk['text'])
                
                st.write("**Metadata:**")
                st.json(chunk['metadata'])

# Footer
st.divider()
st.markdown("""
<div style='text-align: center; color: #666; margin-top: 2rem;'>
    <p>🚀 GTI SOP Assistant - Unified Version</p>
    <p>All-in-one document processing, search, and chat interface</p>
</div>
""", unsafe_allow_html=True)
